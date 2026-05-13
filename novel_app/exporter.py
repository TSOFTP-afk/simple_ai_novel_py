from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from novel_app.database import Database

try:
    from docx import Document as _DocxDocument

    HAS_DOCX = True
except ImportError:  # pragma: no cover - optional runtime dependency guard
    _DocxDocument = None
    HAS_DOCX = False


class BookExporter:
    def __init__(self, database: Database) -> None:
        self.database = database

    def export_txt(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        output_path.write_text(self._build_book_txt(payload), encoding="utf-8")
        return output_path

    def export_markdown(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        output_path.write_text(self._build_book_markdown(payload), encoding="utf-8")
        return output_path

    def export_docx(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        document = self._build_book_docx(payload)
        document.save(str(output_path))
        return output_path

    def export_chapter_txt(self, chapter_id: int, output_path: Path) -> Path:
        payload = self._get_chapter_payload(chapter_id)
        output_path.write_text(self._build_chapter_txt(payload), encoding="utf-8")
        return output_path

    def export_chapter_markdown(self, chapter_id: int, output_path: Path) -> Path:
        payload = self._get_chapter_payload(chapter_id)
        output_path.write_text(self._build_chapter_markdown(payload), encoding="utf-8")
        return output_path

    def export_chapter_docx(self, chapter_id: int, output_path: Path) -> Path:
        payload = self._get_chapter_payload(chapter_id)
        document = self._new_docx()
        self._append_chapter_docx(document, payload["chapter"], heading_level=1)
        document.save(str(output_path))
        return output_path

    def _get_chapter_payload(self, chapter_id: int) -> dict[str, Any]:
        chapter = self.database.get_chapter(chapter_id)
        if not chapter:
            raise ValueError("Chapter not found.")
        book = self.database.get_book(int(chapter["book_id"]))
        return {
            "book": dict(book) if book else {"title": ""},
            "chapter": dict(chapter),
        }

    def _build_book_txt(self, payload: dict[str, Any]) -> str:
        book = payload["book"]
        groups = self._group_chapters(payload["chapters"])
        use_volume_headings = self._use_volume_headings(payload)

        lines = [f"书名：{book['title']}", "=" * 40, ""]
        if book.get("outline_text"):
            lines.extend(["书籍大纲：", str(book["outline_text"]).strip(), ""])

        if payload["characters"]:
            lines.extend(["人物卡：", "-" * 20])
            for item in payload["characters"]:
                lines.append(f"{item['name']} / {item['role'] or '未设定角色'}")
                if item["profile_text"]:
                    lines.append(item["profile_text"])
                lines.append("")

        if payload["world_entries"]:
            lines.extend(["世界观：", "-" * 20])
            for item in payload["world_entries"]:
                lines.append(f"{item['name']} [{item['category']}]")
                if item["content_text"]:
                    lines.append(item["content_text"])
                lines.append("")

        lines.extend(["-" * 40, ""])
        for group in groups:
            if use_volume_headings:
                lines.append(f"【{group['title']}】")
                lines.append("")
            for chapter in group["chapters"]:
                lines.extend(self._chapter_txt_lines(chapter))

        return "\n".join(lines).strip() + "\n"

    def _build_book_markdown(self, payload: dict[str, Any]) -> str:
        book = payload["book"]
        groups = self._group_chapters(payload["chapters"])
        use_volume_headings = self._use_volume_headings(payload)

        lines = [f"# {book['title']}", ""]
        if book.get("outline_text"):
            lines.extend(["## 书籍大纲", "", str(book["outline_text"]).strip(), ""])

        if payload["characters"]:
            lines.extend(["## 人物卡", ""])
            for item in payload["characters"]:
                lines.append(f"### {item['name']}")
                lines.append("")
                lines.append(f"- 角色：{item['role'] or '未设定'}")
                if item["profile_text"]:
                    lines.append("")
                    lines.append(item["profile_text"])
                lines.append("")

        if payload["world_entries"]:
            lines.extend(["## 世界观", ""])
            for item in payload["world_entries"]:
                lines.append(f"### {item['name']}")
                lines.append("")
                lines.append(f"- 分类：{item['category']}")
                lines.append("")
                if item["content_text"]:
                    lines.append(item["content_text"])
                    lines.append("")

        for group in groups:
            if use_volume_headings:
                lines.append(f"### {group['title']}")
                lines.append("")
            for chapter in group["chapters"]:
                lines.extend(self._chapter_markdown_lines(chapter, heading_level=4))

        return "\n".join(lines).strip() + "\n"

    def _build_book_docx(self, payload: dict[str, Any]):
        document = self._new_docx()
        book = payload["book"]
        document.add_heading(str(book["title"]), level=0)
        if book.get("outline_text"):
            document.add_heading("书籍大纲", level=1)
            self._append_multiline_paragraphs(document, str(book["outline_text"]))

        if payload["characters"]:
            document.add_heading("人物卡", level=1)
            for item in payload["characters"]:
                document.add_heading(str(item["name"]), level=2)
                document.add_paragraph(f"角色：{item['role'] or '未设定'}")
                self._append_multiline_paragraphs(document, str(item["profile_text"] or ""))

        if payload["world_entries"]:
            document.add_heading("世界观", level=1)
            for item in payload["world_entries"]:
                document.add_heading(str(item["name"]), level=2)
                document.add_paragraph(f"分类：{item['category']}")
                self._append_multiline_paragraphs(document, str(item["content_text"] or ""))

        groups = self._group_chapters(payload["chapters"])
        use_volume_headings = self._use_volume_headings(payload)
        for group in groups:
            if use_volume_headings:
                document.add_heading(str(group["title"]), level=2)
            for chapter in group["chapters"]:
                self._append_chapter_docx(document, chapter, heading_level=3)
        return document

    def _build_chapter_txt(self, payload: dict[str, Any]) -> str:
        book = payload["book"]
        chapter = payload["chapter"]
        lines = [f"书名：{book.get('title', '')}", f"章节：{chapter['title']}", "=" * 40, ""]
        lines.extend(self._chapter_txt_lines(chapter, include_title=False))
        return "\n".join(lines).strip() + "\n"

    def _build_chapter_markdown(self, payload: dict[str, Any]) -> str:
        book = payload["book"]
        chapter = payload["chapter"]
        lines = [f"# {chapter['title']}", "", f"> 书籍：{book.get('title', '')}", ""]
        lines.extend(self._chapter_markdown_lines(chapter, include_title=False))
        return "\n".join(lines).strip() + "\n"

    def _chapter_txt_lines(self, chapter: dict[str, Any], include_title: bool = True) -> list[str]:
        lines: list[str] = []
        if include_title:
            lines.extend([f"第{chapter['sort_order']}章 {chapter['title']}", ""])
        if chapter.get("outline"):
            lines.extend(["大纲：", str(chapter["outline"]).strip(), ""])
        lines.extend([str(chapter.get("content") or "").strip(), ""])
        return lines

    def _chapter_markdown_lines(
        self,
        chapter: dict[str, Any],
        *,
        heading_level: int = 2,
        include_title: bool = True,
    ) -> list[str]:
        lines: list[str] = []
        if include_title:
            prefix = "#" * max(1, min(6, heading_level))
            lines.extend([f"{prefix} 第{chapter['sort_order']}章 {chapter['title']}", ""])
        if chapter.get("outline"):
            lines.extend(["### 大纲", "", str(chapter["outline"]).strip(), ""])
        lines.append(str(chapter.get("content") or "").strip())
        lines.append("")
        return lines

    def _append_chapter_docx(self, document, chapter: dict[str, Any], heading_level: int = 2) -> None:
        title = str(chapter.get("title") or "未命名章节")
        sort_order = chapter.get("sort_order")
        heading = f"第{sort_order}章 {title}" if sort_order else title
        document.add_heading(heading, level=max(1, min(4, heading_level)))
        if chapter.get("outline"):
            document.add_heading("大纲", level=max(1, min(5, heading_level + 1)))
            self._append_multiline_paragraphs(document, str(chapter["outline"]))
        self._append_multiline_paragraphs(document, str(chapter.get("content") or ""))

    def _append_multiline_paragraphs(self, document, text: str) -> None:
        paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
        if not paragraphs:
            return
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    def _new_docx(self):
        if not HAS_DOCX or _DocxDocument is None:
            raise RuntimeError("当前环境缺少 python-docx，无法导出 DOCX。")
        return _DocxDocument()

    def _group_chapters(self, chapters: list[dict[str, Any]]) -> list[dict[str, Any]]:
        groups: list[dict[str, Any]] = []
        current_title: str | None = None

        for chapter in chapters:
            group_title = chapter["volume_title"] or "未分卷"
            if group_title != current_title:
                groups.append({"title": group_title, "chapters": []})
                current_title = group_title
            groups[-1]["chapters"].append(chapter)

        return groups

    def _use_volume_headings(self, payload: dict[str, Any]) -> bool:
        if payload["volumes"]:
            return True
        return any(chapter["volume_id"] is not None for chapter in payload["chapters"])

    def export_content_only_txt(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        groups = self._group_chapters(payload["chapters"])
        lines: list[str] = []
        for group in groups:
            for chapter in group["chapters"]:
                content = str(chapter.get("content") or "").strip()
                if content:
                    lines.append(content)
                    lines.append("")
        output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        return output_path

    def export_outline_only_txt(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        groups = self._group_chapters(payload["chapters"])
        lines = [f"# {payload['book']['title']} 大纲", ""]
        for group in groups:
            for chapter in group["chapters"]:
                lines.append(f"## 第{chapter['sort_order']}章 {chapter['title']}")
                outline = str(chapter.get("outline") or "").strip()
                if outline:
                    lines.append(outline)
                lines.append("")
        output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        return output_path

    def export_characters_only_txt(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        lines = [f"# {payload['book']['title']} 人物卡", ""]
        for item in payload["characters"]:
            lines.append(f"## {item['name']} / {item['role'] or '未设定角色'}")
            if item["profile_text"]:
                lines.append(item["profile_text"])
            lines.append("")
        output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        return output_path

    def export_world_only_txt(self, book_id: int, output_path: Path) -> Path:
        payload = self.database.get_book_export_data(book_id)
        lines = [f"# {payload['book']['title']} 世界观", ""]
        for item in payload["world_entries"]:
            lines.append(f"## {item['name']} [{item['category']}]")
            if item["content_text"]:
                lines.append(item["content_text"])
            lines.append("")
        output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        return output_path

    def export_per_volume_txt(self, book_id: int, output_path: Path) -> list[Path]:
        payload = self.database.get_book_export_data(book_id)
        groups = self._group_chapters(payload["chapters"])
        paths: list[Path] = []
        base_name = output_path.stem
        parent_dir = output_path.parent
        for group in groups:
            safe_name = re.sub(r"[\\/:*?\"<>|]", "_", group["title"])
            vol_path = parent_dir / f"{base_name}_{safe_name}.txt"
            lines = [f"书名：{payload['book']['title']}", f"# {group['title']}", "=" * 40, ""]
            for chapter in group["chapters"]:
                lines.extend(self._chapter_txt_lines(chapter))
            vol_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
            paths.append(vol_path)
        return paths
